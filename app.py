import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# --- ページ設定 ---
st.set_page_config(page_title="CHICAGO Setlist Builder", layout="wide")

# --- データの読み込み（今回はエラー防止のためサンプルデータを使用） ---
@st.cache_data
def load_data():
    # ※スプレッドシート連携の代わりに、まずは動くサンプルデータを入れます
    data = {
        "楽曲名": ["25 or 6 to 4", "Saturday in the Park", "Hard to Say I'm Sorry", "If You Leave Me Now", "Does Anybody Really Know What Time It Is?", "Beginnings", "Make Me Smile", "Colour My World", "Just You 'n' Me", "Feelin' Stronger Every Day"],
        "Key": ["Am", "C", "E", "B", "C", "A", "E", "F", "C", "D"],
        "Harp (Pos)": ["", "C", "", "", "", "A", "", "", "", ""],
        "BPM": ["150", "114", "72", "100", "112", "120", "130", "60", "90", "120"],
        "備考・演奏指示": ["", "ピアノから", "バラード", "", "ホーンセクション", "アコギ", "", "フルート", "", "アップテンポ"]
    }
    return pd.DataFrame(data)

# --- Wordドキュメント作成関数 ---
def create_docx(setlist_df, main_count):
    doc = Document()
    
    # タイトル
    title = doc.add_heading('LIVE SETLIST', 0)
    title.alignment = 1 # 中央揃え
    
    # 表の作成（1行目はヘッダー）
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '順'
    hdr_cells[1].text = '楽曲名'
    hdr_cells[2].text = 'Key / Harp'
    hdr_cells[3].text = '備考'
    
    # 選んだ曲を表に追加していく
    for i, row in setlist_df.iterrows():
        row_cells = table.add_row().cells
        
        # 予備曲の判定（指定した本番曲数を超えたら「予備」にする）
        if i >= main_count:
            row_cells[0].text = "予備"
        else:
            row_cells[0].text = str(i + 1)
            
        row_cells[1].text = str(row['楽曲名'])
        row_cells[2].text = f"Key: {row.get('Key', '')}\nHarp: {row.get('Harp (Pos)', '')}"
        row_cells[3].text = str(row.get('備考・演奏指示', ''))

    # メモリ上にWordデータを保存
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# --- メイン画面構築 ---
st.title("🎸 CHICAGO セトリ作成システム")

# 状態の初期化
if 'setlist' not in st.session_state:
    st.session_state.setlist = pd.DataFrame(columns=["順番", "楽曲名", "Key", "Harp (Pos)", "BPM", "備考・演奏指示"])

df_master = load_data()

# タブの作成
tab1, tab2, tab3 = st.tabs(["🔍 1. 楽曲検索・追加", "📝 2. セトリ編集・曲順入替", "📄 3. 出力プレビュー"])

# --- タブ1: 楽曲の検索と追加 ---
with tab1:
    st.markdown("### マスターデータから曲を選ぶ")
    st.caption("左端のチェックボックスにチェックを入れた曲が、セトリに追加されます。")
    
    # チェックボックス付きでマスターデータを表示
    df_selectable = df_master.copy()
    df_selectable.insert(0, "追加", False)
    
    edited_df = st.data_editor(
        df_selectable,
        hide_index=True,
        column_config={"追加": st.column_config.CheckboxColumn("追加", default=False)},
        use_container_width=True
    )
    
    if st.button("➕ チェックした曲をセトリに追加"):
        selected_rows = edited_df[edited_df["追加"] == True].drop(columns=["追加"])
        if not selected_rows.empty:
            # 現在のセトリの最後尾の順番を取得
            current_max_order = st.session_state.setlist["順番"].max()
            if pd.isna(current_max_order):
                current_max_order = 0
                
            # 追加する曲に順番を割り振る
            selected_rows.insert(0, "順番", range(int(current_max_order) + 1, int(current_max_order) + 1 + len(selected_rows)))
            
            # セトリに結合
            st.session_state.setlist = pd.concat([st.session_state.setlist, selected_rows], ignore_index=True)
            st.success(f"{len(selected_rows)}曲を追加しました！「セトリ編集」タブを確認してください。")
        else:
            st.warning("曲が選択されていません。")

# --- タブ2: セトリの編集（曲順入れ替え・キー変更） ---
with tab2:
    st.markdown("### 今日のセットリストを編集")
    st.caption("💡 **「順番」の数字を書き換える**と、その順番通りに並び替わります！Keyや備考もここで直接修正できます。")
    
    if st.session_state.setlist.empty:
        st.info("まだ曲が追加されていません。「1. 楽曲検索・追加」タブから曲を選んでください。")
    else:
        # 順番を入れ替えたり、内容を編集できるデータエディター
        edited_setlist = st.data_editor(
            st.session_state.setlist,
            num_rows="dynamic", # 行の削除も可能にする
            use_container_width=True,
            hide_index=True
        )
        
        # 編集された内容を保存し、順番（数値）で並び替える
        if not edited_setlist.equals(st.session_state.setlist):
            # 順番カラムを数値に変換してソート
            edited_setlist["順番"] = pd.to_numeric(edited_setlist["順番"], errors='coerce').fillna(999)
            edited_setlist = edited_setlist.sort_values(by="順番").reset_index(drop=True)
            # 順番を1から振り直す
            edited_setlist["順番"] = range(1, len(edited_setlist) + 1)
            
            st.session_state.setlist = edited_setlist
            st.rerun() # 画面を更新して並び替えを反映

        if st.button("🗑️ セトリをすべてクリア", type="secondary"):
            st.session_state.setlist = pd.DataFrame(columns=["順番", "楽曲名", "Key", "Harp (Pos)", "BPM", "備考・演奏指示"])
            st.rerun()

# --- タブ3: プレビューとダウンロード ---
with tab3:
    st.markdown("### 出力プレビューと設定")
    
    if st.session_state.setlist.empty:
        st.info("曲が追加されていません。")
    else:
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("#### ⚙️ 曲数設定")
            total_songs = len(st.session_state.setlist)
            st.write(f"現在の登録数: **{total_songs}曲**")
            
            # 本番曲数の設定スライダー
            main_count = st.number_input("本番の曲数（これ以降は「予備」になります）", min_value=1, max_value=total_songs, value=total_songs)
            
            st.markdown("#### 📥 ダウンロード")
            st.caption("ボタンを押すとWordファイル(.docx)が保存されます。")
            
            # Wordデータの生成
            docx_data = create_docx(st.session_state.setlist, main_count)
            
            st.download_button(
                label="📥 Word形式でダウンロード",
                data=docx_data,
                file_name="CHICAGO_Setlist.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
            
        with col2:
            st.markdown("#### 📄 プレビュー (Wordのイメージ)")
            
            # プレビュー用の表示（Wordの中身を模して表示）
            preview_df = st.session_state.setlist.copy()
            
            # 予備曲のラベル付け
            preview_labels = []
            for i in range(len(preview_df)):
                if i >= main_count:
                    preview_labels.append("🛑 予備")
                else:
                    preview_labels.append(f" {i+1} ")
            
            preview_df.insert(0, "出力表示", preview_labels)
            
            st.dataframe(
                preview_df[["出力表示", "楽曲名", "Key", "Harp (Pos)", "備考・演奏指示"]],
                hide_index=True,
                use_container_width=True
            )
